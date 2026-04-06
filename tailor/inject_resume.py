"""
tailor/inject_resume.py
-----------------------
Injects tailored content from tailor_resume.py draft into your
resume .docx template, producing a ready-to-review application copy.

Architecture (variable-length sections):
  - The template has ONE section placeholder per experience block,
    e.g. {{GRA_SECTION}}, {{IBM_SR_SECTION}}, etc.
  - The tailoring draft contains exactly the bullets Claude decided
    to write for each section — 2, 3, or 4 as appropriate.
  - This script reads those bullets, counts them, and writes exactly
    that many paragraphs into the .docx — cloning formatting from
    the placeholder paragraph. No manual --cuts needed.

First-time setup (run once):
  python tailor/inject_resume.py --setup

Normal usage:
  python tailor/inject_resume.py \\
    --draft "data/tailored/Qualcomm_ML_Engineer_2025-04-05.txt" \\
    --company "Qualcomm" --role "ML Engineer Tools" \\
    --location cleveland

  python tailor/inject_resume.py \\
    --draft "data/tailored/Amazon_Data_Scientist_2025-04-05.txt" \\
    --company "Amazon" --role "Data Scientist" \\
    --location relocate
"""

import os
import re
import sys
import logging
import argparse
import platform
import subprocess
from copy import deepcopy
from datetime import date, datetime
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from dotenv import load_dotenv

load_dotenv()

# ----------------------------------------------------------------
# PATHS
# ----------------------------------------------------------------
BASE_DIR      = Path(__file__).resolve().parent.parent
DATA_DIR      = BASE_DIR / "data"
TAILORED_DIR  = DATA_DIR / "tailored"
TEMPLATE_PATH = BASE_DIR / "resume_template.docx"

TAILORED_DIR.mkdir(parents=True, exist_ok=True)

# ----------------------------------------------------------------
# LOGGING
# ----------------------------------------------------------------
today_str = date.today().isoformat()
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s  %(levelname)-8s  %(message)s",
    handlers=[logging.StreamHandler()],
)
log = logging.getLogger(__name__)

# ----------------------------------------------------------------
# LOCATION MAP
# ----------------------------------------------------------------
LOCATION_MAP = {
    "cleveland":  "Cleveland, OH",
    "relocate":   "Open to Relocate",
    "remote":     "Remote",
    "houston":    "Houston, TX",
    "seattle":    "Seattle, WA",
    "sf":         "San Francisco, CA",
    "nyc":        "New York, NY",
    "boston":     "Boston, MA",
    "austin":     "Austin, TX",
    "chicago":    "Chicago, IL",
    "sandiego":   "San Diego, CA",
    "sanjose":    "San Jose, CA",
}

# ----------------------------------------------------------------
# SECTION PLACEHOLDER TAGS
# One tag per section in the template.
# ----------------------------------------------------------------
SECTION_TAGS = {
    "{{LOCATION}}":          "location",
    "{{SUMMARY}}":           "summary",
    "{{GRA_SECTION}}":       "gra",
    "{{IBM_SR_SECTION}}":    "ibm_sr",
    "{{IBM_ASSOC_SECTION}}": "ibm_assoc",
    "{{PROJECT_SECTION}}":   "projects",
    "{{SKILLS_SECTION}}":    "skills",
}

# ----------------------------------------------------------------
# DRAFT PARSER
# ----------------------------------------------------------------

class DraftParser:
    """
    Parses the structured .txt output of tailor_resume.py.

    Looks for section headers matching the tailoring_system_prompt.txt
    output format, then extracts REWRITTEN: blocks under each.
    """

    SECTION_HEADERS = {
        "summary":   ["SUMMARY"],
        "gra":       ["GRA BULLETS", "GRA BULLET"],
        "ibm_sr":    ["IBM SR BULLETS", "IBM SR BULLET", "IBM SENIOR BULLETS"],
        "ibm_assoc": ["IBM ASSOC BULLETS", "IBM ASSOC BULLET", "IBM ASSOCIATE BULLETS"],
        "projects":  ["PROJECT BULLETS", "PROJECT BULLET"],
        "skills":    ["SKILLS LINES", "SKILLS LINE"],
    }

    def __init__(self, draft_path: Path):
        self.text = draft_path.read_text(encoding="utf-8")
        self._sections: dict[str, list[str]] = {k: [] for k in self.SECTION_HEADERS}
        self._omit: set[str] = set()
        self._parse()

    def _parse(self) -> None:
        # The draft file is divided by ======= separators into chunks.
        # SECTION 1 (tailoring) is the chunk immediately after the
        # "SECTION 1 — RESUME BULLET TAILORING" label chunk.
        # We split on the long === lines and find the right chunk.
        chunks = re.split(r"={40,}", self.text)

        body = ""
        for i, chunk in enumerate(chunks):
            if "SECTION 1" in chunk and "RESUME BULLET TAILORING" in chunk:
                # The content is in the NEXT chunk
                if i + 1 < len(chunks):
                    body = chunks[i + 1]
                break

        if not body:
            # Fallback: try to find content between section markers directly
            m = re.search(
                r"SECTION 1.*?TAILORING.*?={40,}(.*?)(?:={40,}|SECTION 2|\Z)",
                self.text,
                re.DOTALL | re.IGNORECASE,
            )
            body = m.group(1) if m else self.text
            log.warning("Used fallback parser — section structure may differ from expected")

        if not body.strip():
            log.error("Draft parser found empty body — no bullets will be injected")
            return

        # Split on "---" dividers within the body
        blocks = re.split(r"\n-{2,}\s*\n", body)

        current_key = None
        for block in blocks:
            block = block.strip()
            if not block:
                continue
            matched = self._match_header(block)
            if matched:
                current_key = matched
                # Content may follow the header in the same block
                content = re.sub(r"^[A-Z][A-Z &]+\n?", "", block).strip()
                if content and current_key:
                    self._extract(current_key, content)
            elif current_key:
                self._extract(current_key, block)

        self._log()

    def _match_header(self, text: str) -> str | None:
        first = text.split("\n")[0].strip().upper()
        for key, headers in self.SECTION_HEADERS.items():
            for h in headers:
                if h in first:
                    return key
        return None

    def _extract(self, key: str, text: str) -> None:
        if "OMIT_SECTION" in text.upper():
            self._omit.add(key)
            return
        bullets = re.findall(
            r"REWRITTEN:\s*\n(.*?)(?=\nREWRITTEN:|\nORIGINAL:|\nCHANGE NOTES:|\Z)",
            text,
            re.DOTALL,
        )
        for b in bullets:
            c = b.strip()
            if c:
                self._sections[key].append(c)

    def _log(self) -> None:
        parts = []
        for key in self.SECTION_HEADERS:
            if key in self._omit:
                parts.append(f"{key}:OMIT")
            else:
                parts.append(f"{key}:{len(self._sections[key])}")
        log.info("Draft parsed — " + " | ".join(parts))

    def get(self, section: str) -> list[str]:
        return self._sections.get(section, [])

    def is_omitted(self, section: str) -> bool:
        return section in self._omit


# ----------------------------------------------------------------
# PARAGRAPH CLONING — format-preserving
# ----------------------------------------------------------------

def clone_para_with_text(source_para, new_text: str):
    """
    Deep-copy a paragraph's XML and replace text content,
    preserving style, numbering, indentation, and run formatting.
    """
    new_p = deepcopy(source_para._element)
    runs  = new_p.findall(f".//{qn('w:r')}")

    if not runs:
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = new_text
        r.append(t)
        new_p.append(r)
        return new_p

    # Keep first run's formatting, remove the rest
    first_run = runs[0]
    for run in runs[1:]:
        run.getparent().remove(run)

    t_elem = first_run.find(qn("w:t"))
    if t_elem is None:
        t_elem = OxmlElement("w:t")
        first_run.append(t_elem)

    t_elem.text = new_text
    if new_text.startswith(" ") or new_text.endswith(" "):
        t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

    return new_p


def clean_bullet_text(text: str) -> str:
    """
    Post-process bullet text from the tailoring draft:
      1. Strip markdown bold markers (**text**)
      2. Replace em-dash variants (" —", "—") with ","
    """
    # Strip markdown bold: **some text** → some text
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    # Replace " —" and "—" with ","
    text = text.replace(" —", ",").replace("—", ",")
    return text.strip()


def replace_para_text(para, new_text: str) -> None:
    """Replace single paragraph text, preserving first-run formatting."""
    if not para.runs:
        para.text = new_text
        return
    first_rpr = deepcopy(para.runs[0]._r.find(qn("w:rPr")))
    for run in list(para.runs):
        run._r.getparent().remove(run._r)
    new_run = para.add_run(new_text)
    if first_rpr is not None:
        new_run._r.insert(0, first_rpr)


def add_bottom_border_to_para(para) -> None:
    """
    Add a thin bottom border to a paragraph (used for section headings
    to recreate the visual separator line).
    """
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        para._element.insert(0, pPr)

    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)

    # Remove existing bottom if present
    existing = pBdr.find(qn('w:bottom'))
    if existing is not None:
        pBdr.remove(existing)

    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')        # 0.5pt — thin line
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)


def replace_location_run_only(para, location_value: str) -> None:
    """
    Replace ONLY the {{LOCATION}} token in the first plain run of the
    location paragraph, leaving all hyperlink runs untouched.
    This fixes issue 3 (doubled contact info).
    """
    for elem in para._element.iter():
        if elem.tag == qn('w:t') and '{{LOCATION}}' in (elem.text or ''):
            elem.text = (elem.text or '').replace('{{LOCATION}}', location_value)
            if elem.text.startswith(' ') or elem.text.endswith(' '):
                elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            return


def make_skills_run_bold_prefix(para, text: str) -> None:
    """
    Write a skills line with the prefix (up to and including the colon)
    in bold and the rest in normal weight.
    Inherits paragraph-level font (rFonts, sz) so the style is preserved.
    """
    colon_idx = text.find(':')
    if colon_idx == -1:
        replace_para_text(para, text)
        return

    bold_part   = text[:colon_idx + 1] + ' '
    normal_part = text[colon_idx + 2:]

    # Get paragraph-level rPr for font/size
    pPr = para._element.find(qn('w:pPr'))
    para_rpr = deepcopy(pPr.find(qn('w:rPr'))) if pPr is not None else None

    # Clear existing runs
    for run in list(para.runs):
        run._r.getparent().remove(run._r)

    def run_with_rpr(text_val: str, bold: bool):
        r = OxmlElement('w:r')
        rpr = deepcopy(para_rpr) if para_rpr is not None else OxmlElement('w:rPr')
        # Remove existing bold
        for tag in [qn('w:b'), qn('w:bCs')]:
            e = rpr.find(tag)
            if e is not None:
                rpr.remove(e)
        if bold:
            rpr.insert(0, OxmlElement('w:bCs'))
            rpr.insert(0, OxmlElement('w:b'))
        else:
            b_off = OxmlElement('w:b')
            b_off.set(qn('w:val'), '0')
            rpr.insert(0, b_off)
        r.append(rpr)
        t = OxmlElement('w:t')
        t.text = text_val
        if text_val.startswith(' ') or text_val.endswith(' '):
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r.append(t)
        return r

    para._element.append(run_with_rpr(bold_part, bold=True))
    para._element.append(run_with_rpr(normal_part, bold=False))


# Project hyperlink lookup
# Key: substring that appears in the project title (case-insensitive)
# Value: GitHub URL
PROJECT_URLS: dict[str, str] = {
    "shakespearean":         "https://github.com/Aisha-AI-Dev/Shakespearean_LLM",
    "blendedrag":            "https://github.com/cwru-courses/csds497-f24-2/tree/main/Project/axp1343",
    "enhanced retrieval":    "https://github.com/cwru-courses/csds497-f24-2/tree/main/Project/axp1343",
    "rag pipelines":         "https://github.com/cwru-courses/csds497-f24-2/tree/main/Project/axp1343",
    "arithmetic reasoning":  "https://github.com/Aisha-AI-Dev/LLM",
    "compact transformers":  "https://github.com/Aisha-AI-Dev/LLM",
    "hybrid movie":          "https://github.com/Aisha-AI-Dev/Recommendation-System",
    "recommender":           "https://github.com/Aisha-AI-Dev/Recommendation-System",
    "hebbian":               "https://github.com/Aisha-AI-Dev/Biologically-Plausible-Neural-Network-Training",
    "biologically":          "https://github.com/Aisha-AI-Dev/Biologically-Plausible-Neural-Network-Training",
    "contrastive hebbian":   "https://github.com/Aisha-AI-Dev/Biologically-Plausible-Neural-Network-Training",
    # P3, P6, P7 have no links — omit intentionally
}


def make_project_run_with_hyperlink(para, text: str) -> None:
    """
    Write a project bullet with the project title as a hyperlink.
    Expected format: "Project Name: description text"
    The title (up to colon) becomes bold + hyperlinked if URL exists.
    Falls back to bold-only if no URL is configured.

    Font fix: copies paragraph-level rPr into each run so the
    resume's custom style font (minorHAnsi theme) is respected.
    """
    colon_idx = text.find(':')
    if colon_idx == -1:
        replace_para_text(para, text)
        return

    title       = text[:colon_idx].strip()
    description = text[colon_idx + 1:].strip()

    # Get the paragraph-level rPr for font/size inheritance
    pPr = para._element.find(qn('w:pPr'))
    para_rpr = deepcopy(pPr.find(qn('w:rPr'))) if pPr is not None else None

    def make_rpr(bold: bool, hyperlink: bool = False) -> object:
        """Build a w:rPr element inheriting paragraph font."""
        rpr = deepcopy(para_rpr) if para_rpr is not None else OxmlElement('w:rPr')
        # Remove existing bold elements before setting
        for tag in [qn('w:b'), qn('w:bCs'), qn('w:rStyle')]:
            existing = rpr.find(tag)
            if existing is not None:
                rpr.remove(existing)
        if hyperlink:
            rStyle = OxmlElement('w:rStyle')
            rStyle.set(qn('w:val'), 'Hyperlink')
            rpr.insert(0, rStyle)
        if bold:
            b = OxmlElement('w:b')
            bcs = OxmlElement('w:bCs')
            rpr.insert(0, bcs)
            rpr.insert(0, b)
        else:
            # Explicitly turn off bold
            b = OxmlElement('w:b')
            b.set(qn('w:val'), '0')
            rpr.insert(0, b)
        return rpr

    # Clear existing runs
    for run in list(para.runs):
        run._r.getparent().remove(run._r)
    # Clear any leftover hyperlink elements too
    for h in para._element.findall(qn('w:hyperlink')):
        para._element.remove(h)

    # Find URL for this project
    url = None
    title_lower = title.lower()
    for key, link in PROJECT_URLS.items():
        if key.lower() in title_lower:
            url = link
            break

    title_text = title + ':'

    if url:
        # Build hyperlink element
        r_id = para.part.relate_to(
            url,
            'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
            is_external=True
        )
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id', r_id)
        hyperlink.set(qn('w:history'), '1')

        r_elem = OxmlElement('w:r')
        r_elem.append(make_rpr(bold=True, hyperlink=True))
        t = OxmlElement('w:t')
        t.text = title_text
        r_elem.append(t)
        hyperlink.append(r_elem)
        para._element.append(hyperlink)
    else:
        # Bold title, no hyperlink
        r_elem = OxmlElement('w:r')
        r_elem.append(make_rpr(bold=True))
        t = OxmlElement('w:t')
        t.text = title_text
        r_elem.append(t)
        para._element.append(r_elem)

    # Normal run for description — single space separator, no trailing space
    r_desc = OxmlElement('w:r')
    r_desc.append(make_rpr(bold=False))
    t_desc = OxmlElement('w:t')
    t_desc.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t_desc.text = ' ' + description.rstrip()
    r_desc.append(t_desc)
    para._element.append(r_desc)




def set_spacing_after(para, twips: int) -> None:
    """Set w:spacing w:after on a paragraph's pPr."""
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        para._element.insert(0, pPr)
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:after'), str(twips))


def expand_section(anchor_para, bullets: list[str]) -> None:
    """
    Replace anchor_para with N cloned paragraphs (one per bullet).
    Sets spacing-after to 0 on the last paragraph so there is no
    visual gap before the next section heading.
    """
    for text in reversed(bullets):
        new_p = clone_para_with_text(anchor_para, text)
        anchor_para._element.addnext(new_p)
    anchor_para._element.getparent().remove(anchor_para._element)


def remove_para(para) -> None:
    """Remove a paragraph entirely from the document."""
    para._element.getparent().remove(para._element)


# ----------------------------------------------------------------
# INJECTION
# ----------------------------------------------------------------

def inject(doc: Document,
           parser: DraftParser,
           location_value: str) -> dict[str, str]:
    """Walk paragraphs, find section tags, inject content."""
    status = {}

    # Take a snapshot of paragraphs at start — we mutate during iteration
    all_paras = list(doc.paragraphs)

    for para_idx, para in enumerate(all_paras):
        text = para.text.strip()
        if not text:
            continue

        for tag, section in SECTION_TAGS.items():
            if tag not in text:
                continue

            if section == "location":
                replace_location_run_only(para, location_value)
                status["location"] = location_value

            elif section == "summary":
                bullets = parser.get("summary")
                if bullets:
                    cleaned = clean_bullet_text(bullets[0])
                    replace_para_text(para, cleaned)
                    status["summary"] = "1 line"
                else:
                    log.warning("No summary in draft — keeping placeholder")
                    status["summary"] = "kept placeholder"

                # Fix 2: add bottom border to the SUMMARY heading (para before this)
                if para_idx > 0:
                    summary_heading = all_paras[para_idx - 1]
                    if summary_heading.text.strip().upper() == "SUMMARY":
                        add_bottom_border_to_para(summary_heading)

            elif section == "skills":
                bullets = parser.get("skills")
                if bullets:
                    cleaned_bullets = [clean_bullet_text(b) for b in bullets]
                    # Write first bullet in place with bold prefix
                    make_skills_run_bold_prefix(para, cleaned_bullets[0])
                    # Insert remaining bullets after, in order
                    for skill_text in reversed(cleaned_bullets[1:]):
                        new_p = clone_para_with_text(para, skill_text)
                        para._element.addnext(new_p)
                    # Now find the newly inserted paragraphs and apply bold prefix
                    # They were inserted directly after para in the XML
                    live_paras = list(doc.paragraphs)
                    try:
                        live_idx = next(i for i, p in enumerate(live_paras) if p._element is para._element)
                        for j, skill_text in enumerate(cleaned_bullets[1:]):
                            target_idx = live_idx + 1 + j
                            if target_idx < len(live_paras):
                                make_skills_run_bold_prefix(live_paras[target_idx], skill_text)
                    except StopIteration:
                        pass
                    status["skills"] = f"{len(bullets)} lines"
                    log.info(f"  {tag}: {len(bullets)} skill line(s)")
                else:
                    log.warning(f"  {tag}: no skills in draft — placeholder kept")
                    status["skills"] = "no content"

            elif section == "projects":
                bullets = parser.get("projects")
                if parser.is_omitted(section):
                    remove_para(para)
                    status[section] = "OMITTED"
                elif bullets:
                    cleaned = [clean_bullet_text(b) for b in bullets]
                    # Write first bullet in place
                    make_project_run_with_hyperlink(para, cleaned[0])
                    # Single bullet: keep original after spacing (section gap)
                    # Multiple bullets: zero after on first so no gap between them
                    if len(cleaned) > 1:
                        set_spacing_after(para, 0)
                    # Insert remaining after, in order
                    for proj_text in reversed(cleaned[1:]):
                        new_p = clone_para_with_text(para, proj_text)
                        para._element.addnext(new_p)
                    # Apply hyperlink handler to newly inserted paras.
                    # All mid-section bullets get after=0; last bullet keeps
                    # the original after=100 to preserve the section-end gap.
                    live_paras = list(doc.paragraphs)
                    try:
                        live_idx = next(i for i, p in enumerate(live_paras) if p._element is para._element)
                        last_j = len(cleaned) - 2  # index into cleaned[1:]
                        for j, proj_text in enumerate(cleaned[1:]):
                            target_idx = live_idx + 1 + j
                            if target_idx < len(live_paras):
                                target = live_paras[target_idx]
                                make_project_run_with_hyperlink(target, proj_text)
                                # Only zero spacing on non-last bullets
                                if j < last_j:
                                    set_spacing_after(target, 0)
                                # Last bullet: restore 100 for section-end gap
                                else:
                                    set_spacing_after(target, 100)
                    except StopIteration:
                        pass
                    status["projects"] = f"{len(bullets)} projects"
                    log.info(f"  {tag}: {len(bullets)} project(s)")
                else:
                    log.warning(f"  {tag}: no projects in draft")
                    status["projects"] = "no content"

            else:
                if parser.is_omitted(section):
                    remove_para(para)
                    status[section] = "OMITTED"
                    log.info(f"  {tag}: section omitted")
                else:
                    bullets = parser.get(section)
                    if bullets:
                        cleaned = [clean_bullet_text(b) for b in bullets]
                        expand_section(para, cleaned)
                        status[section] = f"{len(bullets)} bullets"
                        log.info(f"  {tag}: {len(bullets)} bullet(s)")
                    else:
                        log.warning(f"  {tag}: no bullets in draft — placeholder kept")
                        status[section] = "no content"
            break

    return status


# ----------------------------------------------------------------
# SETUP — create template from existing resume
# ----------------------------------------------------------------

def create_template(source_path: Path, output_path: Path) -> None:
    """
    One-time setup: copy the resume and replace section bullets
    with single section-level placeholder tags.
    """
    doc = Document(source_path)

    # First bullet of each section → becomes the section tag
    # All subsequent bullets in the section → deleted
    section_first = {
        "Open to Relocate":                   "{{LOCATION}}",
        "Data scientist with 4+":             "{{SUMMARY}}",
        "Researching how UI/UX":              "{{GRA_SECTION}}",
        "Boosted revenue forecasting":        "{{IBM_SR_SECTION}}",
        "Automated key business":             "{{IBM_ASSOC_SECTION}}",
        "Shakespearean LLM":                  "{{PROJECT_SECTION}}",
        "NLP & GenAI:":                       "{{SKILLS_SECTION}}",
    }

    section_extras = {
        "{{GRA_SECTION}}":        [
            "Engineered an LLM-powered",
            "Co-designed and co-instructed",
            "Investigated hybrid",
        ],
        "{{IBM_SR_SECTION}}":     [
            "Automated clause generation",
            "Developed an in-house sentiment",
            "Led cross-functional finance",
            "Streamlined executive decision",
            "Designed and deployed robust ETL",
            "Collaborated cross-functionally",
        ],
        "{{IBM_ASSOC_SECTION}}":  [
            "Scripted and optimized mainframe",
            "Deployed internal ML models via Flask",
            "Built a real-time alerting pipeline",
            "Reconciled global asset records",
        ],
        "{{PROJECT_SECTION}}":    [
            "Enhanced Retrieval for RAG",
            "IoT Water Quality",
            "Arithmetic Reasoning",
            "Hybrid Movie Recommender",
        ],
        "{{SKILLS_SECTION}}":     [
            "ML & Deep Learning:",
            "Languages:",
            "Cloud & MLOps:",
            "Data Science Techniques:",
        ],
    }

    to_delete = set()
    for extras in section_extras.values():
        to_delete.update(extras)

    changes = 0
    deleted = 0

    for para in list(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        if any(text.startswith(p) for p in to_delete):
            remove_para(para)
            deleted += 1
            continue

        for prefix, tag in section_first.items():
            if text.startswith(prefix):
                replace_para_text(para, tag)
                log.info(f"  Tagged: '{text[:55]}' → {tag}")
                changes += 1
                break

    doc.save(output_path)
    log.info(f"\nTemplate: {changes} tags, {deleted} bullets removed → {output_path}")
    log.info("Open in Word/LibreOffice to verify before using.")


# ----------------------------------------------------------------
# HELPERS
# ----------------------------------------------------------------

def slugify(s: str) -> str:
    return s.strip().replace(" ", "_").replace("/", "-")[:35]


def build_output_path(company: str, role: str) -> Path:
    """
    Returns the output .docx path, creating a company subfolder if needed.

    Structure:
      data/tailored/
        McKinsey/
          McKinsey_Business_Analyst_Tech_AI_2026-04-05_RESUME.docx
        Qualcomm/
          Qualcomm_ML_Engineer_Tools_2026-04-05_RESUME.docx

    - Company folder name is the raw company string, spaces preserved,
      truncated to 50 chars and stripped of path-unsafe characters.
    - If the folder doesn't exist it is created automatically.
    - Filename uses the slugified company + role + date as before.
    """
    # Safe folder name: keep spaces and caps, strip path separators
    safe_company = re.sub(r'[/\\:*?"<>|]', '', company.strip())[:50]
    company_dir  = TAILORED_DIR / safe_company
    company_dir.mkdir(parents=True, exist_ok=True)

    filename = f"{slugify(company)}_{slugify(role)}_{today_str}_RESUME.docx"
    return company_dir / filename


def open_doc(path: Path) -> None:
    try:
        sys_name = platform.system()
        if sys_name == "Darwin":
            subprocess.run(["open", str(path)], check=True)
        elif sys_name == "Windows":
            os.startfile(str(path))
        else:
            for app in ["libreoffice", "xdg-open"]:
                try:
                    subprocess.run([app, str(path)], check=True)
                    break
                except FileNotFoundError:
                    continue
        log.info("Opened for review.")
    except Exception as e:
        log.warning(f"Could not open automatically: {e}")
        log.info(f"Open manually: {path}")


# ----------------------------------------------------------------
# MAIN
# ----------------------------------------------------------------

def main(draft_path: Path, company: str, role: str,
         location_key: str, no_open: bool = False) -> None:

    log.info("=" * 60)
    log.info(f"Resume injector — {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    log.info("=" * 60)

    if not draft_path.exists():
        log.error(f"Draft not found: {draft_path}")
        sys.exit(1)
    if not TEMPLATE_PATH.exists():
        log.error(f"Template not found. Run: python tailor/inject_resume.py --setup")
        sys.exit(1)

    location_value = LOCATION_MAP.get(
        location_key.lower().replace(" ", ""), location_key
    )
    log.info(f"Location: {location_value}")

    parser = DraftParser(draft_path)
    doc    = Document(TEMPLATE_PATH)

    log.info("Injecting...")
    status = inject(doc, parser, location_value)

    log.info("Result:")
    for k, v in status.items():
        log.info(f"  {k:<22} → {v}")

    out = build_output_path(company, role)
    doc.save(out)
    log.info(f"\nSaved → {out}")

    if not no_open:
        open_doc(out)

    log.info("=" * 60)
    log.info("REVIEW: accuracy → bullet count → location → formatting → PDF")
    log.info("=" * 60)


# ----------------------------------------------------------------
# ENTRY POINT
# ----------------------------------------------------------------

if __name__ == "__main__":
    ap = argparse.ArgumentParser(
        description="Inject tailored bullets into resume .docx — bullet count auto-detected from draft",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    ap.add_argument("--setup",    action="store_true",
                    help="One-time: create resume_template.docx from your resume")
    ap.add_argument("--draft",    type=Path, default=None,
                    help="Tailored draft .txt from tailor_resume.py")
    ap.add_argument("--company",  type=str, default="Company")
    ap.add_argument("--role",     type=str, default="Role")
    ap.add_argument("--location", type=str, default="relocate",
                    help="Location key (cleveland/relocate/remote/etc) or city string")
    ap.add_argument("--no-open",  action="store_true",
                    help="Don't open output file automatically")
    ap.add_argument("--source",   type=Path,
                    default=Path("/mnt/user-data/uploads/1775171740734_RESUME_-_Aishani_Patil.docx"))

    args = ap.parse_args()

    if args.setup:
        src = args.source if args.source.exists() else \
              Path("/mnt/user-data/uploads/1775171740734_RESUME_-_Aishani_Patil.docx")
        create_template(src, TEMPLATE_PATH)
    else:
        if not args.draft:
            log.error("--draft is required")
            sys.exit(1)
        main(args.draft, args.company, args.role, args.location, args.no_open)